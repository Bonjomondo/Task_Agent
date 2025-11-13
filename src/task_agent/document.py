"""
Document generation and format conversion utilities
"""

import os
import logging
from pathlib import Path
from typing import Optional
import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Generate documents in various formats"""
    
    def __init__(self, output_dir: str = "output/generated"):
        """
        Initialize document generator
        
        Args:
            output_dir: Directory to save generated documents
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Initialized DocumentGenerator with output_dir: {output_dir}")
    
    def save_markdown(self, content: str, filename: str) -> str:
        """
        Save content as Markdown file
        
        Args:
            content: Markdown content
            filename: Output filename (without extension)
            
        Returns:
            Path to saved file
        """
        filepath = self.output_dir / f"{filename}.md"
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        logger.info(f"Saved Markdown to: {filepath}")
        return str(filepath)
    
    def markdown_to_docx(self, markdown_content: str, filename: str) -> str:
        """
        Convert Markdown to DOCX format
        
        Args:
            markdown_content: Markdown content
            filename: Output filename (without extension)
            
        Returns:
            Path to saved file
        """
        filepath = self.output_dir / f"{filename}.docx"
        
        # Create document
        doc = Document()
        
        # Parse markdown and add to document
        lines = markdown_content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line - add spacing
                doc.add_paragraph()
                continue
            
            # Handle headings
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], level=3)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('#### '):
                heading = doc.add_heading(line[5:], level=4)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Handle lists
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                # Numbered list
                text = line.split('. ', 1)[1] if '. ' in line else line
                doc.add_paragraph(text, style='List Number')
            # Regular paragraph
            else:
                doc.add_paragraph(line)
        
        # Save document
        doc.save(filepath)
        logger.info(f"Saved DOCX to: {filepath}")
        return str(filepath)
    
    def markdown_to_pdf(self, markdown_content: str, filename: str) -> str:
        """
        Convert Markdown to PDF format
        
        Args:
            markdown_content: Markdown content
            filename: Output filename (without extension)
            
        Returns:
            Path to saved file
        """
        filepath = self.output_dir / f"{filename}.pdf"
        
        # Create PDF
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='Justify',
            parent=styles['BodyText'],
            alignment=TA_JUSTIFY
        ))
        
        # Parse markdown and add to PDF
        lines = markdown_content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                elements.append(Spacer(1, 12))
                continue
            
            # Handle headings
            if line.startswith('# '):
                p = Paragraph(line[2:], styles['Heading1'])
                elements.append(p)
                elements.append(Spacer(1, 12))
            elif line.startswith('## '):
                p = Paragraph(line[3:], styles['Heading2'])
                elements.append(p)
                elements.append(Spacer(1, 12))
            elif line.startswith('### '):
                p = Paragraph(line[4:], styles['Heading3'])
                elements.append(p)
                elements.append(Spacer(1, 12))
            # Handle lists
            elif line.startswith('- ') or line.startswith('* '):
                p = Paragraph(f"• {line[2:]}", styles['BodyText'])
                elements.append(p)
            elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                text = line.split('. ', 1)[1] if '. ' in line else line
                p = Paragraph(text, styles['BodyText'])
                elements.append(p)
            # Regular paragraph
            else:
                p = Paragraph(line, styles['Justify'])
                elements.append(p)
                elements.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(elements)
        logger.info(f"Saved PDF to: {filepath}")
        return str(filepath)
    
    def generate_all_formats(self, markdown_content: str, filename: str) -> dict:
        """
        Generate document in all formats (MD, DOCX, PDF)
        
        Args:
            markdown_content: Markdown content
            filename: Base filename (without extension)
            
        Returns:
            Dictionary with paths to all generated files
        """
        result = {
            'md': self.save_markdown(markdown_content, filename),
            'docx': self.markdown_to_docx(markdown_content, filename),
            'pdf': self.markdown_to_pdf(markdown_content, filename)
        }
        
        logger.info(f"Generated all formats for: {filename}")
        return result
